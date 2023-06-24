from flask import Flask, request, jsonify,Response
import sys
import ast
from experiments.ask_a_question_togpt import ask_a_question
from werkzeug.utils import secure_filename
sys.path.append('H:\myprojects\medical_softwares')
import docExtractors
from docx import Document
import GPTs
app = Flask(__name__)
import os
ALLOWED_EXTENSIONS = {'docx'}
app.config['UPLOAD_FOLDER'] = 'static'

@app.route('/api/uploadDoc', methods=['POST'])
def uploadDoc():
    if 'file' not in request.files:
        return 'No file uploaded', 400
    file = request.files['file']
    # Check if the file has a valid filename
    if file.filename == '':
        return 'Empty filename', 400
    # Check if the file extension is allowed
    if not allowed_file(file.filename):
        return 'Invalid file extension', 400
    # Save the file with a secure filename to the UPLOAD_FOLDER
    filename = secure_filename(file.filename)
    file.save(os.path.join(app.config['UPLOAD_FOLDER'], "SRS.pdf"))
    return 'File uploaded successfully'

@app.route('/api/updateDoc', methods=['GET', 'POST'])
def updateDoc():
    input_file = "SRS.docx"
    doc = Document(input_file)
    updated_doc = Document()
    # Iterate over the elements in the document and print their index and content
    para = []
    tbls = []
    paragraphs = doc.paragraphs
    tables = doc.tables
    para_counter = 0
    tbl_counter = 0
    text_data = ''
    for index, element in enumerate(doc.element.body):
        if "p" in str(element):
            para.append(element)
            para_text = paragraphs[para_counter].text + "\n"
            updated_doc.add_paragraph(para_text)
            text_data += para_text
            para_counter = para_counter + 1
        elif "tbl" in str(element):
            tbls.append(element)
            table_data = docExtractors.read_table(tables[tbl_counter])
            # table_text = tabulate_table(table_data)
            output_from_gpt = ask_a_question(f"table_text: {str(table_data)}")
            if "response:" in output_from_gpt:
                updated_table = output_from_gpt.split("response:")[1].strip()
                try:
                    updated_table = ast.literal_eval(updated_table)
                    table = docExtractors.create_table(updated_doc, updated_table)
                except Exception as e:
                    updated_table = output_from_gpt
                    updated_doc.add_paragraph(updated_table)
                    print("Error in typecasting due to: ", e)
            elif "table_text:" in output_from_gpt:
                updated_table = output_from_gpt.split("table_text:")[1].strip()
                try:
                    updated_table = ast.literal_eval(updated_table)
                    table = docExtractors.create_table(updated_doc, updated_table)
                except Exception as e:
                    updated_table = output_from_gpt
                    updated_doc.add_paragraph(updated_table)
                    print("Error in typecasting due to: ", e)
            else:
                updated_table = output_from_gpt
                updated_doc.add_paragraph(updated_table)
            print(updated_table)

            text_data += str(table_data) + "\n"
            tbl_counter = tbl_counter + 1

    updated_doc.save("output_document.docx")
    open("word_text.txt", 'w', encoding='utf-8').write(text_data)
    print(text_data)
    return jsonify({"updated_doc": "Doc has been updated successfully! Download Updated Doc Now","updated_data": text_data})





def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS



if __name__ == '__main__':
    app.run(debug=True)


