from docx import Document
from docExtractors.extract_tables import read_table
from openAIModules.deficiencies_mapper import get_deficiencies


def extractDoc(filepath):
    input_file = filepath
    doc = Document(input_file)
    # Iterate over the elements in the document and print their index and content
    para = []
    tbls = []
    paragraphs = doc.paragraphs
    tables = doc.tables
    para_counter = 0
    tbl_counter = 0
    text_data = ''
    for index, element in enumerate(doc.element.body):
        if "p" in str(element):
            para.append(element)
            para_text = paragraphs[para_counter].text + "\n"
            text_data += para_text
            para_counter = para_counter + 1
        elif "tbl" in str(element):
            tbls.append(element)
            table_data = read_table(tables[tbl_counter])
            text_data += str(table_data) + "\n"
            tbl_counter = tbl_counter + 1
    open("../word_text.txt", 'w', encoding='utf-8').write(text_data)
    return text_data

if __name__ == "__main__":
    Doctext = extractDoc("SRS.docx")
    output = get_deficiencies(Doctext)
    print(output)
    open("issues.txt","w",encoding="utf-8").write(output)


