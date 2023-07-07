import docExtractors
from docx import Document
import os
from openAIModules.docAgent import agent
import openai
import logging
import ast
import re
from colorama import Fore, Style
from docx.shared import RGBColor
# Create a custom formatter
class ColorFormatter(logging.Formatter):
    def format(self, record):
        message = super().format(record)
        if record.levelno >= logging.WARNING:
            return f'{Fore.RED}{message}{Style.RESET_ALL}'
        return message

class DocAI:
    '''Takes docx as input and call openai to modify text and save the doc after parsing the openai response
    '''
    log = logging.getLogger(__name__)
    handler = logging.StreamHandler()
    handler.setFormatter(ColorFormatter())
    log.addHandler(handler)
    def GptResponse(self,table_data):
        try:
            output_from_gpt = agent(f"table_text: {str(table_data)}")
            return output_from_gpt,True
        except openai.error.ServiceUnavailableError as e:
            logging.warning("ERROR IN GENERATING RESPONSE.OPENAI SERVER IS UNAVAILABLE. Trying again.......")
            response = self.GptResponse(table_data) #TODO: refector the code to handle time limit error in generating final response
            if response[1]:
                return response
        except openai.error.InvalidRequestError as e:
            logging.warning("ERROR IN GENERATING RESPONSE DUE TO TOKEN LIMIT. REDUCING CONTEXT LENGTH.......")
            print(f"Extracted table data is {table_data}. Table is too large....ending the process. ")
            open("./static/error.docx","w").write(str(e))
            return "ending the task",False
    def ReduceTokens(self):
        pass
    def saveLogs(self):
        '''TODO: This function will save the text generated from openai or any error and present it to web
                        #open("Logs.txt", "a+").write(output_from_gpt)'''
        pass
    def updateDoc(self):
        input_file = "SRS.docx"
        try:
            os.remove("./static/output_document.docx")

        except FileNotFoundError as e:
            print("Finding error.docx...............")
            try:
                os.remove("./static/error.docx")
                print("Deleting error.docx")
            except FileNotFoundError as e:
                pass
            pass
        doc = Document(input_file)
        paragraphs = doc.paragraphs
        tables = doc.tables
        # Iterate over the elements in the document and print their index and content
        para_counter = 0
        tbl_counter = 0
        text_data = ''
        for index, element in enumerate(doc.element.body):
            if "p" in str(element):
                pass #TODO: paragraphs are set as it is, need to modify the structure in future
            elif "tbl" in str(element):
                print("\n=======================================================\n")
                print(f"Entering in table: {tbl_counter}")
                table_data = docExtractors.read_table(tables[tbl_counter])
                self.log.warning("Getting new table data from the SRS DOC....")
                print(f"Inserting table data into OPENAI.......")
                print(f"Getting output from OPENAI......")
                output_from_gpt = self.GptResponse(table_data)[0]
                if output_from_gpt == "ending the task":
                    print("Task Ended.")
                    return "Task ended"
                print(f"Successfully Got the output from OPENAI.")
                pattern = r'\[\[.*'
                matches = re.findall(pattern, output_from_gpt, re.DOTALL)
                if len(matches) != 0:
                    raw_tables = matches[0].replace("'", "'''").strip().split("\n\n")
                    self.log.warning(
                        f"There can be multiple output tables so extract each raw table, here is the length of raw tables: {len(raw_tables)}")
                    for raw_table in raw_tables:
                        updated_table = str(raw_table).replace("'", "'''")
                        print(f"Typecasting table using ast......")
                        try:
                            updated_table = ast.literal_eval(updated_table)
                            docInstance = DocMaster()
                            docInstance.updateTable(tables[tbl_counter], updated_table)
                            # table = docExtractors.create_table(updated_doc, updated_table)
                            self.log.warning("Created table using ast.........")
                        except Exception as e:
                            self.log.warning(f"Error in typecasting due to: {e}\n Skipping table update")
                else:
                    pass
                text_data += str(table_data) + "\n"
                tbl_counter = tbl_counter + 1
        print("Preparing final document to save...........")
        doc.save("./static/output_document.docx")
        open("word_text.txt", 'w', encoding='utf-8').write(text_data)
        print("Wooo!Document Saved Successfully.")



class DocMaster:
    def _get_row(self,table):
        for row in table.rows:
            yield row
    def _get_cells(self,row):
        for index, cell in enumerate(row.cells):
            yield cell

    def updateTable(self,existing_table, new_table_data: list):
        parent_rows = self._get_row(existing_table)
        for row in new_table_data:
            try:
                parent_row_test = next(parent_rows)
                cells = self._get_cells(parent_row_test)
            except:
                # create new row if no more row in parent table
                new_row = existing_table.add_row()
                cells = new_row.cells
                pass
            for index, cell_data in enumerate(row):
                try:
                    cell = next(cells)
                    if cell.text != cell_data:
                        cell.text = cell_data
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                font = run.font
                                font.color.rgb = RGBColor(0x00, 0x80, 0x00)  # Green color
                    else:
                        cell.text = cell_data
                except:
                    cell = cells[index]
                    cell.text = cell_data
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.color.rgb = RGBColor(0x00, 0x80, 0x00)  # Green color



