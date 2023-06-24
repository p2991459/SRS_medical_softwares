from docx import Document
from docx.shared import Inches

import ast
import docx
from tabulate import tabulate
from experiments.ask_a_question_togpt import ask_a_question
import time
start_time = time.time()
def create_table(document,data):
    # Create a table with the same number of rows and columns as the data
    num_rows = len(data)
    num_cols = len(data[0])
    table = document.add_table(rows=num_rows, cols=num_cols)

    # Update the content of the table with the data
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.text = data[i][j]

    return table
def convert_to_docx(input_doc, output_docx):
    # Open the .doc file using python-docx
    doc = Document(input_doc)
    print(doc.__dict__)
    # Save the document as .docx
    doc.save(output_docx)
def read_table(table):
    table_content = []
    for row in table.rows:
        row_content = []
        for cell in row.cells:
            cell_text = cell.text.strip()  # Retrieve the text from the cell and remove leading/trailing spaces
            row_content.append(cell_text)
        table_content.append(row_content)
        # print(table_content)
    return table_content
def tabulate_table(table_content):
    table_text = tabulate(table_content, headers="firstrow", tablefmt="pipe")
    return table_text

# Usage example
input_file = 'SRS.docx'
output_file = 'output.docx'
# convert_to_docx(input_file, output_file)
doc = Document(input_file)
updated_doc = Document()
# Iterate over the elements in the document and print their index and content
para = []
tbls = []
paragraphs = doc.paragraphs
tables  = doc.tables
para_counter = 0
tbl_counter = 0
text_data = ''
for index,element in enumerate(doc.element.body):
    if "p" in str(element):
        para.append(element)
        para_text = paragraphs[para_counter].text + "\n"
        updated_doc.add_paragraph(para_text)
        text_data += para_text
        para_counter = para_counter + 1
    elif "tbl" in str(element):
        tbls.append(element)
        table_data = read_table(tables[tbl_counter])
        # table_text = tabulate_table(table_data)
        output_from_gpt = ask_a_question(f"table_text: {str(table_data)}")
        if "response:" in output_from_gpt:
            updated_table = output_from_gpt.split("response:")[1].strip()
            try:
                updated_table = ast.literal_eval(updated_table)
                table = create_table(updated_doc,updated_table)
            except Exception as e:
                updated_table = output_from_gpt
                updated_doc.add_paragraph(updated_table)
                print("Error in typecasting due to: ",e)
        elif "table_text:" in output_from_gpt:
            updated_table = output_from_gpt.split("table_text:")[1].strip()
            try:
                updated_table = ast.literal_eval(updated_table)
                table = create_table(updated_doc, updated_table)
            except Exception as e:
                updated_table = output_from_gpt
                updated_doc.add_paragraph(updated_table)
                print("Error in typecasting due to: ",e)
        else:
            updated_table = output_from_gpt
            updated_doc.add_paragraph(updated_table)
        print(updated_table)

        text_data += str(table_data) + "\n"
        tbl_counter = tbl_counter + 1

updated_doc.save("output_document.docx")
open("word_text.txt",'w',encoding='utf-8').write(text_data)
print(text_data)
# for paragraph in paragraphs:
#     print(paragraph.text)


