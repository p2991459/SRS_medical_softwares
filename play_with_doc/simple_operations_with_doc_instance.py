import docx

def get_paragraph_index(doc):
  """Get the index of each paragraph in a doc instance.

  Args:
    doc: The doc instance.

  Returns:
    A list of paragraph indexes.
  """

  paragraph_index = []
  for paragraph in doc.paragraphs:
    paragraph_index.append(paragraph._element.index())

  return paragraph_index


if __name__ == "__main__":
  doc = docx.Document("SRS.docx")
  for index, paragraph in enumerate(doc.paragraphs):
      print(f"Paragraph Index: {index}, Content: {paragraph.text}")