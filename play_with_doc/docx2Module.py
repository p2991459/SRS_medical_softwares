import docx

doc = docx.Document()
doc.add_heading('GeeksForGeeks', 0)
para = doc.add_paragraph(
    '''GeeksforGeeks is a Computer Science portal for geeks.''')
para.add_run(
    ''' It contains well written, well thought and well-explained ''').italic = True

# Adding more content to paragraph
para.add_run('''computer science and programming articles, quizzes etc.''')

# Now save the document to a location
doc.save('gfg.docx')