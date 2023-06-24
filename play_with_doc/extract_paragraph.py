from docx import Document

def convert_to_docx(input_doc, output_docx):
    # Open the .doc file using python-docx
    doc = Document(input_doc)
    print(doc.__dict__)
    # Save the document as .docx
    doc.save(output_docx)


# Usage example
input_file = 'SRS.docx'
output_file = 'output.docx'
# convert_to_docx(input_file, output_file)
doc = Document(input_file)

for paragraph in doc.paragraphs:
    print(paragraph.text)

# print(doc.element.body.xml) //To get the doc body as xml


'''As an expert in evaluating Software Requirement Specifications (SRS) in the medical field using  IEC 62304 standard, your role is to identify a few deficiencies and areas where improvements can be made to align with the IEC 62304 standards'''

''' improve above Software requirement specification based on IEC62304 standards from above listed deficiencies.'''